from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TARGET_DIR = Path(r"D:\仿真花独立站项目交接包_20260430")
DOCX_PATH = TARGET_DIR / "01-项目交接记录.docx"
README_PATH = TARGET_DIR / "00-打开说明.txt"
FONT = "Microsoft YaHei"


def set_font(run, size=10.5, bold=False, color=RGBColor(32, 32, 29)):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_font(
        run,
        size=15 if level == 1 else 12,
        bold=True,
        color=RGBColor(57, 67, 51) if level == 1 else RGBColor(32, 32, 29),
    )


def add_body(doc, text, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run("- " + item)
        set_font(run, size=10.2)


def add_path_item(doc, label, path, status):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}：")
    set_font(r1, bold=True)
    r2 = p.add_run(path)
    set_font(r2, size=10.2, color=RGBColor(42, 78, 104))
    r3 = p.add_run(f"  ({status})")
    set_font(r3, size=10, color=RGBColor(111, 106, 97))


def build_docx():
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("仿真花外贸独立站项目交接记录")
    set_font(tr, size=21, bold=True, color=RGBColor(57, 67, 51))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("用于下次继续开发：测试原型、文案表、素材整理和后续建站方向")
    set_font(sr, size=10.5, color=RGBColor(111, 106, 97))

    callout = add_body(
        doc,
        "当前状态：项目已经完成一个可本地预览的静态测试站，英文文案和中英文对照表已整理；"
        "下一步建议先做页面视觉检查，再整理 D 盘素材。",
        space_after=10,
    )
    shade_paragraph(callout, "F0EFE8")

    add_heading(doc, "一、项目目标")
    add_body(
        doc,
        "做一个仿真花外贸独立站测试原型，风格参考 AFLORAL：干净、高级、家居审美强，"
        "但同时突出自家工厂、批发、OEM、定制、样品、包装和出口能力。当前品牌名 Floravelle 只是临时占位，"
        "Logo 和正式品牌名后续都可以替换。",
    )

    add_heading(doc, "二、已经完成")
    add_bullets(
        doc,
        [
            "静态网站原型：已做 6 个页面，包括首页、分类页、产品详情模板、Custom & OEM、Factory、Contact。",
            "英文页面文案：已优化成外贸 B2B 口吻，重点突出工厂直供、批发、样品、定制和包装。",
            "中英文对照文案表：已生成 Excel / CSV / Markdown，方便看不懂英文时对照审核。",
            "产品资料模板：已生成 Excel，用于后续录入产品名称、尺寸、材质、颜色、用途、包装等。",
            "素材要求文档：已整理过素材拍摄要求；当前文件路径未在本次检查中确认到，需要后续按实际文件位置补充。",
        ],
    )

    add_heading(doc, "三、重要文件位置")
    add_path_item(doc, "网站项目目录", r"D:\Codex-Workspace\Projects\floravelle-site", "已确认存在")
    add_path_item(doc, "本地预览入口", r"http://127.0.0.1:4173/index.html", "服务运行时可打开")
    add_path_item(doc, "中英文文案 Excel", r"D:\Codex-Workspace\Projects\floravelle-site\outputs\floravelle-site-copy-bilingual.xlsx", "已确认存在")
    add_path_item(doc, "产品资料模板 Excel", r"D:\Codex-Workspace\Projects\floravelle-site\outputs\floravelle-product-intake-template.xlsx", "已确认存在")
    add_path_item(doc, "中文文案说明", r"D:\Codex-Workspace\Projects\floravelle-site\docs\content\site-copy-bilingual.zh-CN.md", "已确认存在")
    add_path_item(doc, "产品分类说明", r"D:\Codex-Workspace\Projects\floravelle-site\docs\content\product-taxonomy.zh-CN.md", "已确认存在")
    add_path_item(doc, "素材收集指南", r"D:\Codex-Workspace\Projects\floravelle-site\docs\content\asset-collection-guide.zh-CN.md", "已确认存在")
    add_path_item(doc, "用户素材目录", r"D:\花", "约定路径，本次检查未检测到")
    add_path_item(doc, "素材要求 Word", r"D:\外贸独立站资料\仿真花网站素材要求.docx", "历史生成路径，本次检查未检测到")

    add_heading(doc, "四、当前验证结果")
    add_bullets(
        doc,
        [
            "网站检查通过：Site check passed: 6 pages, 3 shared assets.",
            "中英文文案表检查通过：Copy workbook verification passed.",
            "页面里之前出现的乱码分隔符已清理。",
            "当前没有复制整个网站项目，只做交接记录和索引，避免 D 盘重复占空间。",
        ],
    )

    add_heading(doc, "五、未完成事项")
    add_bullets(
        doc,
        [
            "页面视觉检查：需要用电脑和手机宽度检查文字、按钮、表单、导航是否错位。",
            "素材盘点：需要整理 D:\\花 里的图片，判断哪些适合作主图、细节图、场景图，哪些暂不适合上站。",
            "产品分类临时版：需要根据真实产品进一步确认分类结构。",
            "真实图片接入：先选可用图片替换部分占位图，后续再精修。",
            "品牌名和 Logo：当前还未确定，Floravelle 只是临时占位。",
            "正式建站方向：后续再决定 Shopify、WordPress/WooCommerce，或继续静态原型。",
        ],
    )

    add_heading(doc, "六、下一步建议")
    add_bullets(
        doc,
        [
            "第一步：刷新本地网站，做首页、分类页、产品页、定制页、工厂页、联系页的视觉 QA。",
            "第二步：把后续图片统一放到 D:\\花 或另一个固定素材目录，并保持原图不直接修改。",
            "第三步：做素材盘点表，给每张图标注产品类型、可用位置、缺失信息和是否需要补拍。",
            "第四步：从现有素材里选几张最能用的图，替换网站部分占位图，看真实产品放进去的效果。",
            "第五步：再决定临时分类、品牌方向和正式建站平台。",
        ],
    )

    add_heading(doc, "七、项目约束和偏好")
    add_bullets(
        doc,
        [
            "不能误删或侵犯与项目无关的文件。",
            "原始图片只读不改；如果要处理图片，应另存为处理版。",
            "需要下载工具、插件或 skill 时先申请，并优先下载/安装到 D 盘。",
            "暂时不开子代理；后续如果任务足够大、用子代理划算，需要先申请确认。",
            "用户英文阅读不方便，所以重要英文文案要配中文对照说明。",
        ],
    )

    add_heading(doc, "八、明天继续时可以直接这样说")
    add_body(
        doc,
        "“继续仿真花独立站项目，先做页面视觉 QA，然后整理 D 盘素材盘点，不要修改原图。”",
    )

    doc.save(DOCX_PATH)


def build_readme():
    text = """仿真花独立站项目交接包

建议先打开：
1. 01-项目交接记录.docx
2. D:\\Codex-Workspace\\Projects\\floravelle-site\\outputs\\floravelle-site-copy-bilingual.xlsx
3. D:\\Codex-Workspace\\Projects\\floravelle-site\\outputs\\floravelle-product-intake-template.xlsx

本地网站预览：
http://127.0.0.1:4173/index.html

明天继续可以直接说：
继续仿真花独立站项目，先做页面视觉 QA，然后整理 D 盘素材盘点，不要修改原图。

注意：
- 不复制整个网站项目，避免 D 盘重复占空间。
- 原图只读不改，处理图另存。
- 当前品牌名 Floravelle 是临时占位。
"""
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    README_PATH.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_readme()
    print(DOCX_PATH)
    print(README_PATH)
