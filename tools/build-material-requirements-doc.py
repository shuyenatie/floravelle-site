from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(r"D:\外贸独立站资料")
OUTPUT_PATH = OUTPUT_DIR / "仿真花网站素材要求.docx"
FONT = "Microsoft YaHei"


def set_run_font(run, size=10.5, bold=False, color=RGBColor(32, 32, 29)):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(
        run,
        size=15.5 if level == 1 else 12.2,
        bold=True,
        color=RGBColor(57, 67, 51) if level == 1 else RGBColor(32, 32, 29),
    )
    return p


def add_body(doc, text, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run("- " + item)
        set_run_font(run, size=10.2)


def add_grouped_list(doc, groups):
    for title, items in groups:
        add_heading(doc, title, level=2)
        add_bullets(doc, items)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)


def build_doc():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("仿真花外贸独立站素材要求")
    set_run_font(r, size=22, bold=True, color=RGBColor(57, 67, 51))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run("适用于产品图片、场景图、工厂图、包装图和产品资料收集")
    set_run_font(sr, size=10.5, color=RGBColor(111, 106, 97))

    add_heading(doc, "一、核心原则")
    add_body(
        doc,
        "不用一次性拍完所有素材。先保证每个准备上网站的产品有清楚主图和基础资料，"
        "再给重点产品补细节图、插瓶图、场景图和包装图。",
    )

    add_heading(doc, "二、每个产品最低必须有")
    add_bullets(
        doc,
        [
            "产品主图 1 张：完整、清晰、背景尽量干净。",
            "尺寸信息：总长度/高度，花头直径或整体宽度。",
            "材质信息：绢布、PU、塑料、Real Touch、铁丝杆等。",
            "颜色信息：主色、可定制颜色。",
            "用途：家居、婚礼、酒店、户外、批发、零售。",
            "是否支持定制：颜色、长度、包装、标签。",
        ],
    )

    add_heading(doc, "三、重点产品最好准备")
    add_grouped_list(
        doc,
        [
            ("产品主图：1-2 张", ["完整入镜、背景干净、产品居中。"]),
            ("细节图：2-4 张", ["花瓣、叶片、枝干、花头连接处、材质质感。"]),
            ("尺寸对比图：1 张", ["手持、靠尺子，或插在花瓶里看比例。"]),
            ("场景图：1-2 张", ["餐桌、客厅、窗边、婚礼、酒店或展厅。"]),
            ("包装图：1 张", ["内包装、外箱、标签、条码或装箱方式。"]),
        ],
    )

    add_heading(doc, "四、图片拍摄标准")
    add_bullets(
        doc,
        [
            "清晰，不糊；不过曝，不太暗。",
            "产品完整，不裁掉花头或枝杆。",
            "背景尽量简单，不要有杂乱桌面、地面、纸箱、塑料袋。",
            "不要有水印、微信聊天截图、无关文字。",
            "不要强滤镜，同一批产品尽量同一光线、同一角度。",
        ],
    )

    add_heading(doc, "五、主图怎么拍")
    add_bullets(
        doc,
        [
            "背景建议白色、浅灰、米白。",
            "产品居中，整枝完整入镜，四周留一点边距。",
            "长枝产品优先拍竖图；盆栽或成品花艺可以拍 4:5 或 1:1。",
            "主图主要用于产品列表和产品详情页，尽量统一风格。",
        ],
    )

    add_heading(doc, "六、细节图拍什么")
    add_bullets(
        doc,
        [
            "花瓣纹理、花瓣厚度、颜色渐变。",
            "叶片纹理、枝干材质、花头连接处。",
            "是否有铁丝可弯曲，Real Touch / PU / 绢布的质感。",
            "成品组合的花量密度。",
        ],
    )

    add_heading(doc, "七、场景图怎么解决")
    add_body(
        doc,
        "如果暂时没有专业场景图，可以先用简单场景拍摄。后续也可以用 AI 场景图或合成图"
        "辅助首页、分类页和氛围图，但产品详情页最好仍以真实产品图为主。",
    )
    add_bullets(
        doc,
        [
            "简单可行场景：干净花瓶、白墙、窗边、餐桌、柜子、样品间。",
            "背景少放杂物，光线自然，可以只拍局部，不一定拍完整房间。",
            "AI 或合成图适合首页和分类氛围图，不建议完全替代产品真实图。",
        ],
    )

    add_heading(doc, "八、工厂 / OEM 素材")
    add_bullets(
        doc,
        [
            "样品间、产品陈列墙、打样过程。",
            "生产过程、质检过程、包装过程。",
            "外箱、标签、吊牌、条码、仓库、出货照片。",
            "工厂图要干净、有秩序。太乱、太暗、地面堆货的图先不要放。",
        ],
    )

    add_heading(doc, "九、包装和物流素材")
    add_bullets(
        doc,
        [
            "单个产品包装图、内盒/袋子图、外箱图。",
            "每箱装多少、外箱尺寸、毛重/净重。",
            "是否可定制包装，是否可贴客户标签。",
            "这些信息会影响报价、物流体积重和客户信任。",
        ],
    )

    add_heading(doc, "十、尺寸怎么记录")
    add_grouped_list(
        doc,
        [
            (
                "单枝 / 花枝",
                [
                    "总长度 Total Length。",
                    "花头直径 Bloom Diameter。",
                    "枝杆长度 Stem Length。",
                    "展开宽度 Overall Width。",
                ],
            ),
            (
                "盆栽 / 成品花艺",
                [
                    "总高度 Total Height。",
                    "整体宽度 Overall Width。",
                    "花盆高度 Pot Height。",
                    "花盆直径 Pot Diameter。",
                ],
            ),
            (
                "包装",
                ["单个包装尺寸。", "外箱尺寸。", "每箱数量。", "毛重 / 净重。"],
            ),
        ],
    )

    add_heading(doc, "十一、为什么要尺寸")
    add_bullets(
        doc,
        [
            "客户要判断能不能插进花瓶，比例是否适合空间。",
            "婚礼、活动、酒店和软装客户要计算布置数量。",
            "批发商和零售商要判断是否适合货架和包装。",
            "电商卖家需要尺寸写 Listing。",
            "报价、装箱数量、体积重和物流成本都和尺寸有关。",
        ],
    )

    add_heading(doc, "十二、图片命名方式")
    add_body(doc, "建议格式：品类-产品-颜色-尺寸-图片类型-序号")
    add_bullets(
        doc,
        [
            "hydrangea-white-19in-main-01.jpg",
            "hydrangea-white-19in-detail-petal-01.jpg",
            "orchid-potted-white-28in-scene-01.jpg",
            "greenery-hanging-32in-main-01.jpg",
            "绣球-白色-19寸-主图-01.jpg",
            "兰花盆栽-白色-28寸-场景-01.jpg",
        ],
    )

    add_heading(doc, "十三、暂时不要上网站的图片")
    add_bullets(
        doc,
        [
            "太糊、太暗、过曝、颜色严重失真。",
            "背景非常乱，产品被裁掉。",
            "有水印、无关文字、微信聊天截图。",
            "地上随手拍，或者只拍包装袋看不清产品。",
            "看起来明显廉价的工厂堆货图。",
        ],
    )

    add_heading(doc, "十四、当前最应该做的事")
    add_bullets(
        doc,
        [
            "先选 20 个最想卖、最好看、利润也可以的产品。",
            "每个产品至少拍 1 张清楚主图。",
            "记录尺寸、材质、颜色、用途、是否可定制。",
            "从里面挑 5 个重点款补细节图和插瓶图。",
            r"后续素材放在 D:\花，原图不直接修改。",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "素材收集速查清单")
    add_body(doc, "普通产品必须先满足：")
    add_bullets(doc, ["清楚主图。", "总长度/高度。", "材质/颜色。", "用途。", "是否可定制。"])
    add_body(doc, "重点产品额外建议准备：")
    add_bullets(
        doc,
        [
            "细节图 2-4 张。",
            "插瓶或尺寸对比图 1 张。",
            "场景图 1-2 张。",
            "包装图 1 张。",
        ],
    )
    add_body(doc, "工厂 / OEM 素材：")
    add_bullets(doc, ["不需要每款都有，适合重点系列或定制项目。"])

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
